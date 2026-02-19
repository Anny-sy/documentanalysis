"""Legal document processor for handling various document formats."""

import re
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field

from rich.console import Console
from tqdm import tqdm

console = Console()


@dataclass
class DocumentMetadata:
    """Metadata extracted from legal documents."""
    
    filename: str
    file_type: str
    page_count: int = 0
    case_name: Optional[str] = None
    court: Optional[str] = None
    date: Optional[str] = None
    citation: Optional[str] = None
    parties: list[str] = field(default_factory=list)
    judges: list[str] = field(default_factory=list)


@dataclass
class ProcessedDocument:
    """Represents a processed legal document."""
    
    content: str
    metadata: DocumentMetadata
    sections: list[dict] = field(default_factory=list)


class LegalDocumentProcessor:
    """Process legal documents from various formats."""
    
    # Common legal document patterns
    CASE_CITATION_PATTERN = r'\d+\s+[A-Z][a-z]*\.?\s*\d*d?\s+\d+'
    DATE_PATTERN = r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}'
    SECTION_HEADERS = [
        "OPINION", "BACKGROUND", "FACTS", "ANALYSIS", "DISCUSSION",
        "CONCLUSION", "HOLDING", "JUDGMENT", "ORDER", "DISSENT",
        "CONCURRENCE", "PROCEDURAL HISTORY", "STANDARD OF REVIEW",
        "LEGAL STANDARD", "INTRODUCTION", "SUMMARY"
    ]
    
    def __init__(self):
        self._pdf_loader = None
        self._docx_loader = None
    
    def process_file(self, file_path: str | Path) -> ProcessedDocument:
        """Process a single document file."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix == ".pdf":
            return self._process_pdf(file_path)
        elif suffix == ".docx":
            return self._process_docx(file_path)
        elif suffix == ".txt":
            return self._process_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    def process_directory(self, directory: str | Path) -> list[ProcessedDocument]:
        """Process all supported documents in a directory."""
        directory = Path(directory)
        documents = []
        
        supported_extensions = {".pdf", ".docx", ".txt"}
        files = [f for f in directory.rglob("*") if f.suffix.lower() in supported_extensions]
        
        console.print(f"[bold green]Found {len(files)} documents to process[/bold green]")
        
        for file_path in tqdm(files, desc="Processing documents"):
            try:
                doc = self.process_file(file_path)
                documents.append(doc)
            except Exception as e:
                console.print(f"[red]Error processing {file_path}: {e}[/red]")
        
        return documents
    
    def _process_pdf(self, file_path: Path) -> ProcessedDocument:
        """Process a PDF document."""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber is required. Install with: pip install pdfplumber")
        
        text_content = []
        page_count = 0
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
        
        full_text = "\n\n".join(text_content)
        
        metadata = self._extract_metadata(full_text, file_path, "pdf", page_count)
        sections = self._extract_sections(full_text)
        
        return ProcessedDocument(
            content=full_text,
            metadata=metadata,
            sections=sections
        )
    
    def _process_docx(self, file_path: Path) -> ProcessedDocument:
        """Process a DOCX document."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        
        metadata = self._extract_metadata(full_text, file_path, "docx", len(paragraphs))
        sections = self._extract_sections(full_text)
        
        return ProcessedDocument(
            content=full_text,
            metadata=metadata,
            sections=sections
        )
    
    def _process_txt(self, file_path: Path) -> ProcessedDocument:
        """Process a plain text document."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            full_text = f.read()
        
        lines = full_text.split("\n")
        metadata = self._extract_metadata(full_text, file_path, "txt", len(lines))
        sections = self._extract_sections(full_text)
        
        return ProcessedDocument(
            content=full_text,
            metadata=metadata,
            sections=sections
        )
    
    def _extract_metadata(
        self, 
        text: str, 
        file_path: Path, 
        file_type: str,
        page_count: int
    ) -> DocumentMetadata:
        """Extract metadata from document text."""
        metadata = DocumentMetadata(
            filename=file_path.name,
            file_type=file_type,
            page_count=page_count
        )
        
        # Extract citation
        citation_match = re.search(self.CASE_CITATION_PATTERN, text)
        if citation_match:
            metadata.citation = citation_match.group()
        
        # Extract date
        date_match = re.search(self.DATE_PATTERN, text)
        if date_match:
            metadata.date = date_match.group()
        
        # Try to extract case name (usually at the beginning)
        first_lines = text[:500]
        vs_pattern = r'([A-Z][A-Za-z\s,\.]+)\s+v\.?\s+([A-Z][A-Za-z\s,\.]+)'
        case_match = re.search(vs_pattern, first_lines)
        if case_match:
            metadata.case_name = f"{case_match.group(1).strip()} v. {case_match.group(2).strip()}"
            metadata.parties = [case_match.group(1).strip(), case_match.group(2).strip()]
        
        # Extract court information
        court_patterns = [
            r'Supreme Court of [A-Za-z\s]+',
            r'United States Court of Appeals',
            r'United States District Court',
            r'Court of Appeals of [A-Za-z\s]+',
            r'Superior Court of [A-Za-z\s]+'
        ]
        for pattern in court_patterns:
            court_match = re.search(pattern, text[:1000])
            if court_match:
                metadata.court = court_match.group()
                break
        
        return metadata
    
    def _extract_sections(self, text: str) -> list[dict]:
        """Extract logical sections from the document."""
        sections = []
        current_section = {"title": "PREAMBLE", "content": "", "start_pos": 0}
        
        lines = text.split("\n")
        current_pos = 0
        
        for line in lines:
            line_upper = line.strip().upper()
            
            # Check if this line is a section header
            found_header = None
            for header in self.SECTION_HEADERS:
                if line_upper.startswith(header) and len(line_upper) < len(header) + 20:
                    found_header = header
                    break
            
            if found_header:
                # Save previous section
                if current_section["content"].strip():
                    sections.append(current_section)
                
                # Start new section
                current_section = {
                    "title": found_header,
                    "content": "",
                    "start_pos": current_pos
                }
            else:
                current_section["content"] += line + "\n"
            
            current_pos += len(line) + 1
        
        # Add final section
        if current_section["content"].strip():
            sections.append(current_section)
        
        return sections
